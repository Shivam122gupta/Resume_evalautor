"""
document_reader.py
==================
PRESERVED FROM: resumeevaluation/resume_analyzer.py

Original functions: read_pdf, read_docx, read_resume
These are the exact implementations from the working code — only refactored
into a standalone importable module.
"""
from pathlib import Path
from pypdf import PdfReader
from docx import Document


def read_pdf(file_path: str | Path) -> str:
    """Extract text from a PDF file. Preserved from resume_analyzer.py."""
    reader = PdfReader(str(file_path))
    text = ""
    for page in reader.pages:
        page_text = page.extract_text()
        if page_text:
            text += page_text + "\n"
    return text


def read_docx(file_path: str | Path) -> str:
    """Extract text from a DOCX file. Preserved from resume_analyzer.py."""
    document = Document(str(file_path))
    text = ""
    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            text += paragraph.text + "\n"

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += cell.text + "\n"
    return text


def read_resume(file_path: str | Path) -> str | None:
    """
    Dispatcher — reads PDF or DOCX. Preserved from resume_analyzer.py.
    Returns None for unsupported file types.
    """
    path = Path(file_path)
    if path.suffix.lower() == ".pdf":
        return read_pdf(path)
    elif path.suffix.lower() == ".docx":
        return read_docx(path)
    else:
        return None
